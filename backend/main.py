from fastapi import FastAPI, File, UploadFile, HTTPException, Depends, Header
from fastapi.middleware.cors import CORSMiddleware
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials, OAuth2PasswordBearer, OAuth2PasswordRequestForm
from fastapi.responses import FileResponse, JSONResponse
from pydantic import BaseModel, EmailStr, validator
from typing import List, Optional
import uuid
import os
import re
from dotenv import load_dotenv
import fitz  # PyMuPDF
from sentence_transformers import SentenceTransformer
import faiss
import numpy as np
import requests
import json
from docx import Document  # for DOC files
from database import get_db, init_db
from models import User, Document as DBDocument, DocumentChunk, Embedding, SubscriptionType, PlanType
from sqlalchemy.orm import Session
from generate_postman import generate_postman_collection
from datetime import datetime, timedelta
from jose import JWTError, jwt
from passlib.context import CryptContext
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from urllib.parse import quote
from jinja2 import Environment, FileSystemLoader

# Load environment variables
load_dotenv()

# Email Configuration
SMTP_SERVER = os.getenv("SMTP_SERVER", "smtp.gmail.com")
SMTP_PORT = int(os.getenv("SMTP_PORT", "587"))
SMTP_USERNAME = os.getenv("SMTP_USERNAME")
SMTP_PASSWORD = os.getenv("SMTP_PASSWORD")
FRONTEND_URL = os.getenv("FRONTEND_URL", "http://localhost:3000")

# JWT Configuration
SECRET_KEY = os.getenv("SECRET_KEY", "your-secret-key-here")
ALGORITHM = "HS256"
ACCESS_TOKEN_EXPIRE_MINUTES = 30
RESET_TOKEN_EXPIRE_MINUTES = int(os.getenv("RESET_TOKEN_EXPIRE_MINUTES", "15"))
VERIFICATION_TOKEN_EXPIRE_MINUTES = int(os.getenv("VERIFICATION_TOKEN_EXPIRE_MINUTES", "24"))
MAX_LOGIN_ATTEMPTS = 5
LOCKOUT_DURATION_MINUTES = 30

# Password hashing
pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")
oauth2_scheme = OAuth2PasswordBearer(tokenUrl="login")

# Jinja2 setup
template_env = Environment(loader=FileSystemLoader("templates/email"))

# Postman API configuration
POSTMAN_API_KEY = os.getenv("POSTMAN_API_KEY")
POSTMAN_COLLECTION_ID = os.getenv("POSTMAN_COLLECTION_ID")
POSTMAN_API_URL = "https://api.getpostman.com"

# File storage configuration
UPLOAD_DIR = os.getenv("UPLOAD_DIR", "uploaded_docs")
os.makedirs(UPLOAD_DIR, exist_ok=True)

app = FastAPI(
    title="MyQuery API",
    description="Document AI Chatbot API with multi-tenant support",
    version="1.0.0"
)

# Allow CORS for local frontend dev
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

embedding_model = SentenceTransformer("all-MiniLM-L6-v2")
index = None
documents = []

# Define paths for storage
EMBEDDINGS_FILE = "embeddings.json"
DOCUMENTS_FILE = "documents.json"

def send_email(to_email: str, subject: str, template_name: str, template_data: dict):
    """Send email using SMTP with Jinja2 template."""
    try:
        template = template_env.get_template(f"{template_name}.html")
        html_content = template.render(**template_data)
        
        msg = MIMEMultipart()
        msg['From'] = SMTP_USERNAME
        msg['To'] = to_email
        msg['Subject'] = subject

        msg.attach(MIMEText(html_content, 'html'))

        server = smtplib.SMTP(SMTP_SERVER, SMTP_PORT)
        server.starttls()
        server.login(SMTP_USERNAME, SMTP_PASSWORD)
        server.send_message(msg)
        server.quit()
        return True
    except Exception as e:
        print(f"Failed to send email: {str(e)}")
        return False

def create_verification_token(data: dict):
    """Create a token for email verification."""
    to_encode = data.copy()
    expire = datetime.utcnow() + timedelta(minutes=VERIFICATION_TOKEN_EXPIRE_MINUTES)
    to_encode.update({"exp": expire, "type": "verification"})
    encoded_jwt = jwt.encode(to_encode, SECRET_KEY, algorithm=ALGORITHM)
    return encoded_jwt

def create_reset_token(data: dict):
    """Create a token for password reset."""
    to_encode = data.copy()
    expire = datetime.utcnow() + timedelta(minutes=RESET_TOKEN_EXPIRE_MINUTES)
    to_encode.update({"exp": expire, "type": "reset"})
    encoded_jwt = jwt.encode(to_encode, SECRET_KEY, algorithm=ALGORITHM)
    return encoded_jwt

# Security functions
def verify_password(plain_password, hashed_password):
    return pwd_context.verify(plain_password, hashed_password)

def get_password_hash(password):
    return pwd_context.hash(password)

def create_access_token(data: dict):
    to_encode = data.copy()
    expire = datetime.utcnow() + timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES)
    to_encode.update({"exp": expire})
    encoded_jwt = jwt.encode(to_encode, SECRET_KEY, algorithm=ALGORITHM)
    return encoded_jwt

async def get_current_user(token: str = Depends(oauth2_scheme), db: Session = Depends(get_db)):
    credentials_exception = HTTPException(
        status_code=401,
        detail="Could not validate credentials",
        headers={"WWW-Authenticate": "Bearer"},
    )
    try:
        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        email: str = payload.get("sub")
        if email is None:
            raise credentials_exception
    except JWTError:
        raise credentials_exception
    
    user = db.query(User).filter(User.email == email).first()
    if user is None:
        raise credentials_exception
    return user

# Models
class UserCreate(BaseModel):
    name: str
    company_name: str
    email: EmailStr
    password: str
    contact_info: Optional[str] = None
    plan: PlanType
    subscription_type: SubscriptionType

    @validator('password')
    def validate_password_complexity(cls, v):
        if len(v) < 8:
            raise ValueError('Password must be at least 8 characters long')
        if not re.search(r'[A-Z]', v):
            raise ValueError('Password must contain at least one uppercase letter')
        if not re.search(r'[a-z]', v):
            raise ValueError('Password must contain at least one lowercase letter')
        if not re.search(r'\d', v):
            raise ValueError('Password must contain at least one number')
        if not re.search(r'[!@#$%^&*(),.?":{}|<>]', v):
            raise ValueError('Password must contain at least one special character')
        return v

class UserLogin(BaseModel):
    email: EmailStr
    password: str

class Token(BaseModel):
    access_token: str
    token_type: str
    expires_in: int

class UserResponse(BaseModel):
    id: str
    name: str
    company_name: str
    email: str
    plan: PlanType
    subscription_type: SubscriptionType
    subscription_end_date: datetime
    is_active: bool
    is_verified: bool

class UserUpdate(BaseModel):
    name: Optional[str] = None
    company_name: Optional[str] = None
    email: Optional[EmailStr] = None
    contact_info: Optional[str] = None
    plan: Optional[PlanType] = None
    subscription_type: Optional[SubscriptionType] = None
    is_active: Optional[bool] = None

class QueryRequest(BaseModel):
    query: str
    session_id: Optional[str] = None
    user_id: str

# Password reset models
class PasswordResetRequest(BaseModel):
    email: EmailStr

class PasswordReset(BaseModel):
    token: str
    new_password: str
    confirm_password: str

    @validator('confirm_password')
    def passwords_match(cls, v, values, **kwargs):
        if 'new_password' in values and v != values['new_password']:
            raise ValueError('Passwords do not match')
        return v

    @validator('new_password')
    def validate_password_complexity(cls, v):
        if len(v) < 8:
            raise ValueError('Password must be at least 8 characters long')
        if not re.search(r'[A-Z]', v):
            raise ValueError('Password must contain at least one uppercase letter')
        if not re.search(r'[a-z]', v):
            raise ValueError('Password must contain at least one lowercase letter')
        if not re.search(r'\d', v):
            raise ValueError('Password must contain at least one number')
        if not re.search(r'[!@#$%^&*(),.?":{}|<>]', v):
            raise ValueError('Password must contain at least one special character')
        return v

# Helper functions
def calculate_subscription_end_date(start_date: datetime, subscription_type: SubscriptionType) -> datetime:
    if subscription_type == SubscriptionType.MONTHLY:
        return start_date + timedelta(days=30)
    elif subscription_type == SubscriptionType.QUARTERLY:
        return start_date + timedelta(days=90)
    else:  # YEARLY
        return start_date + timedelta(days=365)

# User management endpoints
@app.post("/users", response_model=UserResponse, tags=["User Management"])
async def create_user(user_data: UserCreate, db: Session = Depends(get_db)):
    """Create a new user account with the provided details."""
    # Check if email already exists
    if db.query(User).filter(User.email == user_data.email).first():
        raise HTTPException(status_code=400, detail="Email already registered")
    
    # Calculate subscription end date
    start_date = datetime.utcnow()
    end_date = calculate_subscription_end_date(start_date, user_data.subscription_type)
    
    # Create new user with hashed password
    user = User(
        name=user_data.name,
        company_name=user_data.company_name,
        email=user_data.email,
        password=get_password_hash(user_data.password),
        contact_info=user_data.contact_info,
        plan=user_data.plan,
        subscription_type=user_data.subscription_type,
        subscription_start_date=start_date,
        subscription_end_date=end_date,
        is_verified=False
    )
    db.add(user)
    db.commit()
    db.refresh(user)
    
    # Generate verification token
    verification_token = create_verification_token({"sub": user.email})
    verification_url = f"{FRONTEND_URL}/verify-email?token={quote(verification_token)}"
    
    # Send verification email
    template_data = {
        "name": user.name,
        "verification_url": verification_url,
        "expiry_minutes": VERIFICATION_TOKEN_EXPIRE_MINUTES
    }
    
    if not send_email(user.email, "Verify Your Email Address", "verification", template_data):
        print(f"Failed to send verification email to {user.email}")
    
    return {
        "id": user.id,
        "name": user.name,
        "company_name": user.company_name,
        "email": user.email,
        "plan": user.plan,
        "subscription_type": user.subscription_type,
        "subscription_end_date": user.subscription_end_date,
        "is_active": user.is_active,
        "is_verified": user.is_verified
    }

@app.get("/verify-email", tags=["Authentication"])
async def verify_email(token: str, db: Session = Depends(get_db)):
    """Verify email address using the verification token."""
    try:
        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        email = payload.get("sub")
        token_type = payload.get("type")
        
        if not email or token_type != "verification":
            raise HTTPException(status_code=400, detail="Invalid verification token")
        
        user = db.query(User).filter(User.email == email).first()
        if not user:
            raise HTTPException(status_code=404, detail="User not found")
        
        if user.is_verified:
            return {"message": "Email already verified"}
        
        user.is_verified = True
        db.commit()
        
        return {"message": "Email verified successfully"}
        
    except JWTError:
        raise HTTPException(status_code=400, detail="Invalid or expired verification token")

@app.post("/login", response_model=Token, tags=["Authentication"])
async def login(user_data: UserLogin, db: Session = Depends(get_db)):
    """Login with email and password to get access token."""
    user = db.query(User).filter(User.email == user_data.email).first()
    
    # Check if account is locked
    if user and user.locked_until and user.locked_until > datetime.utcnow():
        remaining_time = (user.locked_until - datetime.utcnow()).total_seconds() / 60
        raise HTTPException(
            status_code=403,
            detail=f"Account is locked. Please try again in {int(remaining_time)} minutes."
        )
    
    if not user or not verify_password(user_data.password, user.password):
        # Increment login attempts
        if user:
            user.login_attempts += 1
            if user.login_attempts >= MAX_LOGIN_ATTEMPTS:
                user.locked_until = datetime.utcnow() + timedelta(minutes=LOCKOUT_DURATION_MINUTES)
            db.commit()
        
        raise HTTPException(
            status_code=401,
            detail="Incorrect email or password",
            headers={"WWW-Authenticate": "Bearer"},
        )
    
    # Reset login attempts on successful login
    user.login_attempts = 0
    user.locked_until = None
    db.commit()
    
    # Check if subscription is active
    if not user.is_active or datetime.utcnow() > user.subscription_end_date:
        raise HTTPException(
            status_code=403,
            detail="Subscription has expired. Please renew your subscription."
        )
    
    access_token = create_access_token(data={"sub": user.email})
    return {
        "access_token": access_token,
        "token_type": "bearer",
        "expires_in": ACCESS_TOKEN_EXPIRE_MINUTES * 60
    }

@app.post("/request-password-reset", tags=["Authentication"])
async def request_password_reset(
    reset_request: PasswordResetRequest,
    db: Session = Depends(get_db)
):
    """Request a password reset token."""
    user = db.query(User).filter(User.email == reset_request.email).first()
    if not user:
        # Don't reveal if email exists or not
        return {"message": "If your email is registered, you will receive a password reset link."}
    
    # Generate reset token
    reset_token = create_reset_token({"sub": user.email})
    reset_url = f"{FRONTEND_URL}/reset-password?token={quote(reset_token)}"
    
    # Send reset email
    template_data = {
        "name": user.name,
        "reset_url": reset_url,
        "expiry_minutes": RESET_TOKEN_EXPIRE_MINUTES
    }
    
    if send_email(user.email, "Password Reset Request", "reset_password", template_data):
        return {"message": "Password reset link has been sent to your email."}
    else:
        raise HTTPException(status_code=500, detail="Failed to send password reset email")

@app.post("/reset-password", tags=["Authentication"])
async def reset_password(
    reset_data: PasswordReset,
    db: Session = Depends(get_db)
):
    """Reset password using the reset token."""
    try:
        payload = jwt.decode(reset_data.token, SECRET_KEY, algorithms=[ALGORITHM])
        email = payload.get("sub")
        token_type = payload.get("type")
        
        if not email or token_type != "reset":
            raise HTTPException(status_code=400, detail="Invalid reset token")
        
        user = db.query(User).filter(User.email == email).first()
        if not user:
            raise HTTPException(status_code=404, detail="User not found")
        
        # Update password
        user.password = get_password_hash(reset_data.new_password)
        user.login_attempts = 0
        user.locked_until = None
        db.commit()
        
        # Send confirmation email
        template_data = {
            "name": user.name
        }
        
        send_email(user.email, "Password Changed", "password_changed", template_data)
        
        return {"message": "Password has been reset successfully"}
        
    except JWTError:
        raise HTTPException(status_code=400, detail="Invalid or expired reset token")

@app.put("/users/{user_id}", response_model=UserResponse, tags=["User Management"])
def update_user(
    user_id: str,
    user_update: UserUpdate,
    db: Session = Depends(get_db)
):
    """Update user information"""
    db_user = db.query(User).filter(User.id == user_id).first()
    if not db_user:
        raise HTTPException(status_code=404, detail="User not found")

    # Update basic fields
    for field, value in user_update.dict(exclude_unset=True).items():
        if field == "subscription_type" and value:
            # Recalculate subscription end date if subscription type changes
            db_user.subscription_end_date = calculate_subscription_end_date(
                db_user.subscription_start_date,
                value
            )
        setattr(db_user, field, value)

    db.commit()
    db.refresh(db_user)
    return db_user

def extract_text_from_pdf(filepath):
    doc = fitz.open(filepath)
    return "\n".join([page.get_text() for page in doc])

def extract_text_from_doc(filepath):
    doc = Document(filepath)
    return "\n".join([paragraph.text for paragraph in doc.paragraphs])

async def process_file(file: UploadFile, user_id: str, db: Session):
    if not file.filename.lower().endswith(('.pdf', '.doc', '.docx')):
        raise HTTPException(status_code=400, detail="Only PDF and DOC files allowed")

    file_id = str(uuid.uuid4())
    filepath = os.path.join(UPLOAD_DIR, f"{file_id}{os.path.splitext(file.filename)[1]}")
    
    content = await file.read()
    with open(filepath, "wb") as f:
        f.write(content)

    # Extract text based on file type
    if file.filename.lower().endswith('.pdf'):
        text = extract_text_from_pdf(filepath)
    else:  # .doc or .docx
        text = extract_text_from_doc(filepath)

    # Create document record
    db_document = DBDocument(
        id=file_id,
        user_id=user_id,
        filename=file.filename,
        content=text
    )
    db.add(db_document)
    db.commit()

    # Split into chunks and create embeddings
    chunk_size = 1000
    chunk_overlap = 100
    chunks = []
    for i in range(0, len(text), chunk_size - chunk_overlap):
        chunk_text = text[i:i+chunk_size]
        chunk = DocumentChunk(
            document_id=file_id,
            content=chunk_text,
            chunk_index=len(chunks)
        )
        db.add(chunk)
        chunks.append(chunk_text)
    
    db.commit()

    # Generate embeddings
    embeddings = embedding_model.encode(chunks)
    
    # Store embeddings
    for chunk, embedding in zip(db.query(DocumentChunk).filter(DocumentChunk.document_id == file_id).all(), embeddings):
        db_embedding = Embedding(
            user_id=user_id,
            chunk_id=chunk.id,
            vector=embedding.tolist()
        )
        db.add(db_embedding)
    
    db.commit()

    # Clean up the file
    os.remove(filepath)
    
    return len(chunks)

@app.post("/upload", tags=["Document Management"])
async def upload_files(
    files: List[UploadFile] = File(...),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Upload one or more documents (PDF or DOC) for processing."""
    if not current_user.is_verified:
        raise HTTPException(
            status_code=403,
            detail="Please verify your email address before uploading documents"
        )
    
    total_chunks = 0
    for file in files:
        chunks = await process_file(file, current_user.id, db)
        total_chunks += chunks

    return {
        "message": f"Processed {len(files)} files",
        "total_chunks": total_chunks
    }

@app.post("/query", tags=["Query"])
async def ask_question(
    query_request: QueryRequest,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Ask a question about the uploaded documents."""
    # Verify user exists
    if current_user.id != query_request.user_id:
        raise HTTPException(status_code=403, detail="Not authorized to query this user's documents")

    # Get all embeddings for the user
    embeddings = db.query(Embedding).filter(Embedding.user_id == query_request.user_id).all()
    if not embeddings:
        raise HTTPException(status_code=400, detail="No documents uploaded for this user")

    # Create FAISS index
    dimension = len(embeddings[0].vector)
    index = faiss.IndexFlatL2(dimension)
    vectors = np.array([e.vector for e in embeddings])
    index.add(vectors)

    # Embed the query
    query_embedding = embedding_model.encode([query_request.query])
    D, I = index.search(np.array(query_embedding), k=3)

    # Get relevant chunks
    chunks = []
    for idx in I[0]:
        chunk = db.query(DocumentChunk).join(Embedding).filter(Embedding.id == embeddings[idx].id).first()
        if chunk:
            chunks.append(chunk.content)

    context = "\n".join(chunks)
    
    # Build conversation history
    history_text = ""
    if query_request.session_id:
        # Implement conversation history storage and retrieval here
        pass

    prompt = f"""You are an empathetic and emotionally intelligent AI assistant. Respond to the following question with both factual accuracy and emotional awareness. 
    Consider the user's potential emotional state and provide a warm, understanding response while maintaining professionalism.

Context from documents:
{context}

{history_text}

Current question: {query_request.query}

Guidelines for your response:
1. Be empathetic and understanding
2. Acknowledge the user's perspective
3. Provide factual information in a warm, engaging way
4. Use appropriate emotional tone based on the question
5. If the question seems emotional or personal, respond with extra care and sensitivity

Answer:"""

    response = requests.post("http://localhost:11434/api/generate", json={"model": "mistral", "prompt": prompt})
    if response.status_code != 200:
        raise HTTPException(status_code=500, detail="LLM failed to generate response")
    
    full_response = ""
    for line in response.text.strip().split('\n'):
        if line:
            try:
                json_response = json.loads(line)
                if 'response' in json_response:
                    full_response += json_response['response']
            except json.JSONDecodeError:
                continue

    return {
        "answer": full_response.strip(),
        "session_id": query_request.session_id or str(uuid.uuid4())
    }

@app.get("/", tags=["Health"])
def read_root():
    """Health check endpoint."""
    return {"message": "Multi-tenant Document AI Chatbot API running"}

@app.get("/generate-postman", tags=["Documentation"])
async def generate_postman_docs():
    """Generate and push the Postman collection directly to Postman."""
    try:
        # Generate the Postman collection
        generate_postman_collection(app)
        
        # Read the generated collection
        with open("myquery.postman_collection.json", "r") as f:
            collection_data = json.load(f)
        
        # Push to Postman API
        headers = {
            "X-Api-Key": POSTMAN_API_KEY,
            "Content-Type": "application/json"
        }
        
        # Update the collection
        response = requests.put(
            f"{POSTMAN_API_URL}/collections/{POSTMAN_COLLECTION_ID}",
            headers=headers,
            json={"collection": collection_data}
        )
        
        if response.status_code not in [200, 201]:
            raise HTTPException(
                status_code=500,
                detail=f"Failed to update Postman collection: {response.text}"
            )
        
        return JSONResponse({
            "message": "Postman collection updated successfully",
            "collection_id": POSTMAN_COLLECTION_ID
        })
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to update Postman collection: {str(e)}")

@app.on_event("startup")
async def startup_event():
    """Initialize database and update Postman collection on startup."""
    # Initialize database
    init_db()
    
    # Generate and update Postman collection
    try:
        generate_postman_collection(app)
        
        # Push to Postman API if API key is configured
        if POSTMAN_API_KEY:
            with open("myquery.postman_collection.json", "r") as f:
                collection_data = json.load(f)
            
            headers = {
                "X-Api-Key": POSTMAN_API_KEY,
                "Content-Type": "application/json"
            }
            
            response = requests.put(
                f"{POSTMAN_API_URL}/collections/{POSTMAN_COLLECTION_ID}",
                headers=headers,
                json={"collection": collection_data}
            )
            
            if response.status_code not in [200, 201]:
                print(f"Warning: Failed to update Postman collection: {response.text}")
            else:
                print("Postman collection updated successfully")
    except Exception as e:
        print(f"Warning: Failed to update Postman collection: {str(e)}")